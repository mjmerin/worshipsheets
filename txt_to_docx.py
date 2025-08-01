from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
import re
import os

def create_two_column_docx(*texts, output_filename="multi_column_output.docx", 
                          font_name='Arial', font_size=11, column_spacing=0.5):
    """
    Create a two-column .docx file from multiple text inputs with automatic page breaks.
    Automatically formats section titles (Verse, Chorus, Bridge, Vamp) in bold with larger font.
    Also makes the first line of each text block bold and 10 points larger.
    
    Args:
        *texts: Variable number of text strings to include
        output_filename (str): Name of the output .docx file
        font_name (str): Font name for the document
        font_size (int): Base font size for the document
        column_spacing (float): Space between columns in inches
    
    Returns:
        str: Path to the created file
    """
    if not texts:
        raise ValueError("At least one text argument must be provided")
    
    doc = Document()
    
    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    
    # Set up two-column layout
    section = doc.sections[0]
    setup_two_columns(section, column_spacing)
    
    # Process each text input
    for i, text in enumerate(texts):
        if i > 0:
            # Add a page break before each new text (except the first)
            add_page_break(doc)
        
        # Add the text content with formatting
        add_formatted_text_to_document(doc, text, font_name, font_size, format_first_line=True)
    
    # Save the document
    doc.save(output_filename)
    print(f"Multi-column document with {len(texts)} text sections saved as: {output_filename}")
    return output_filename

def create_continuous_two_column_docx(*texts, output_filename="lyrics.docx", 
                                     font_name='Arial', font_size=12, column_spacing=0.5, 
                                     separator_text="---"):
    """
    Create a two-column .docx file where all texts flow continuously with formatted titles.
    Also makes the first line of each text block bold and 10 points larger.
    
    Args:
        *texts: Variable number of text strings to include
        output_filename (str): Name of the output .docx file
        font_name (str): Font name for the document
        font_size (int): Base font size for the document
        column_spacing (float): Space between columns in inches
        separator_text (str): Text to separate different input texts
    
    Returns:
        str: Path to the created file
    """
    if not texts:
        raise ValueError("At least one text argument must be provided")
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    
    # Set up two-column layout
    section = doc.sections[0]
    setup_two_columns(section, column_spacing)
    
    # Add all texts with separators
    for i, text in enumerate(texts):
        if i > 0 and separator_text:
            # Add separator between texts
            separator_para = doc.add_paragraph()
            separator_run = separator_para.add_run(separator_text)
            separator_run.font.name = font_name
            separator_run.font.size = Pt(font_size)
            separator_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add the text content with formatting
        add_formatted_text_to_document(doc, text, font_name, font_size, format_first_line=True)
    
    doc.save(output_filename)
    print(f"Continuous two-column document with {len(texts)} text sections saved as: {output_filename}")
    return output_filename

def add_formatted_text_to_document(doc, text, font_name, font_size, format_first_line=False):
    """
    Add text content to document with special formatting for section titles and optionally first line.
    Makes Verse, Chorus, Bridge, Vamp titles bold and 2 points larger.
    If format_first_line is True, makes the first line bold and 10 points larger.
    """
    if not text.strip():
        return
    
    # Split text into lines first to handle first line formatting
    lines = text.split('\n')
    if not lines:
        return
    
    first_line_processed = False
    
    # Process each line
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            # Add empty paragraph for blank lines
            doc.add_paragraph()
            continue
        
        # Check if this line is a section title
        if is_section_title(line):
            # Create paragraph for section title
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.name = font_name
            run.font.size = Pt(font_size + 2)  # 2 points larger
            run.font.bold = True
        elif format_first_line and not first_line_processed and line:
            # Format the first non-empty line specially
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.name = font_name
            run.font.size = Pt(font_size + 10)  # 10 points larger
            run.font.bold = True
            first_line_processed = True
        else:
            # Regular line
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.name = font_name
            run.font.size = Pt(font_size)

def format_first_line_in_docx_text(text, font_name, font_size, doc):
    """
    Alternative approach: Format the first line of a text block and add it to a document.
    This function handles the first line formatting as a separate operation.
    
    Args:
        text (str): The input text block
        font_name (str): Font name to use
        font_size (int): Base font size
        doc (Document): The document to add formatted text to
    
    Returns:
        None (modifies document in place)
    """
    if not text.strip():
        return
    
    lines = text.split('\n')
    first_line = lines[0].strip() if lines else ""
    remaining_lines = lines[1:] if len(lines) > 1 else []
    
    # Add first line with special formatting (unless it's a section title)
    if first_line:
        para = doc.add_paragraph()
        run = para.add_run(first_line)
        run.font.name = font_name
        
        if is_section_title(first_line):
            # Section title formatting
            run.font.size = Pt(font_size + 2)
            run.font.bold = True
        else:
            # First line formatting
            run.font.size = Pt(font_size + 10)
            run.font.bold = True
    
    # Add remaining lines with normal formatting
    for line in remaining_lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()  # Empty paragraph for blank lines
            continue
            
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.name = font_name
        
        if is_section_title(line):
            # Section title formatting
            run.font.size = Pt(font_size + 2)
            run.font.bold = True
        else:
            # Normal formatting
            run.font.size = Pt(font_size)

def is_section_title(text):
    """
    Check if a line of text is a section title (Verse, Chorus, Bridge, Vamp with optional numbers).
    Handles formats like: Verse, Verse1, Verse 1, Chorus, Chorus2, Bridge 3, etc.
    Also handles pipe-delimited formats like |Verse|, |Chorus|, etc.
    """
    # Updated pattern to handle both plain text and pipe-delimited formats
    pattern = r'^(\|)?(Verse|Chorus|Refrain|Pre-Chorus|Pre\s+Chorus|Bridge|Vamp)(\|)?\s*\d*\s*:*$'
    return bool(re.match(pattern, text.strip(), re.IGNORECASE))

def setup_two_columns(section, column_spacing_inches=0.5):
    """Helper function to set up two-column layout for a section"""
    sectPr = section._sectPr
    # Remove existing column settings
    for cols in sectPr.xpath('./w:cols'):
        sectPr.remove(cols)
    
    # Convert inches to twips (1 inch = 1440 twips)
    space_twips = int(column_spacing_inches * 1440)
    
    # Add new column configuration
    cols_xml = f'<w:cols w:num="2" w:space="{space_twips}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    cols = parse_xml(cols_xml)
    sectPr.append(cols)

def add_page_break(doc):
    """Helper function to add a page break"""
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)

def add_column_break(doc):
    """Helper function to add a column break"""
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.COLUMN)

# Example usage
if __name__ == "__main__":
    sample_text1 = """Amazing Grace Song Title
Verse 1
Amazing grace how sweet the sound
That saved a wretch like me
I once was lost but now am found
Was blind but now I see

Chorus
'Twas grace that taught my heart to fear
And grace my fears relieved
How precious did that grace appear
The hour I first believed"""

    sample_text2 = """How Great Thou Art
Verse 1
O Lord my God when I in awesome wonder
Consider all the worlds Thy hands have made
I see the stars, I hear the rolling thunder
Thy power throughout the universe displayed

|Chorus|
Then sings my soul my Savior God to Thee
How great Thou art, how great Thou art"""

    # Test the enhanced functionality
    create_continuous_two_column_docx(
        sample_text1, 
        sample_text2,
        output_filename="test_enhanced_formatting.docx",
        font_size=12
    )